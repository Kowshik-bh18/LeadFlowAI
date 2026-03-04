import io
import logging
from datetime import datetime

from django.shortcuts import get_object_or_404
from django.http import HttpResponse, Http404
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.views import View

logger = logging.getLogger(__name__)


@method_decorator(login_required, name='dispatch')
class ExportDocxView(View):
    def get(self, request, pk):
        from apps.questionnaires.models import Questionnaire
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        questionnaire = get_object_or_404(Questionnaire, pk=pk, user=request.user)
        questions = questionnaire.questions.prefetch_related('answers').order_by('order')

        doc = Document()

        # Title
        title = doc.add_heading('', 0)
        run = title.add_run('LeadFlow AI – Security Questionnaire')
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0xDB)

        doc.add_heading(questionnaire.title, level=1)

        meta = doc.add_paragraph()
        meta.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}').bold = False
        meta.add_run(f'\nTotal Questions: {questionnaire.total_questions}')
        meta.add_run(f'\nAnswered with Citations: {questionnaire.answered_questions}')
        meta.add_run(f'\nNot Found: {questionnaire.not_found_count}')

        doc.add_paragraph()
        doc.add_heading('Questionnaire Responses', level=2)

        for question in questions:
            latest_answer = question.answers.order_by('-created_at').first()

            # Question
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f'Q{question.order}. {question.text}')
            q_run.bold = True
            q_run.font.size = Pt(11)

            if latest_answer:
                # Answer
                a_para = doc.add_paragraph()
                a_para.add_run('Answer: ').bold = True
                a_para.add_run(latest_answer.answer_text)

                # Confidence
                if not latest_answer.is_not_found:
                    conf_para = doc.add_paragraph()
                    conf_run = conf_para.add_run(
                        f'Confidence: {int(latest_answer.confidence_score * 100)}%'
                    )
                    conf_run.italic = True
                    conf_run.font.size = Pt(9)

                # Citations
                if latest_answer.citations:
                    cit_para = doc.add_paragraph()
                    cit_para.add_run('Sources: ').bold = True
                    sources = ', '.join(
                        c.get('document_title', 'Unknown') for c in latest_answer.citations
                    )
                    cit_para.add_run(sources)

            else:
                doc.add_paragraph('Answer: Not yet generated.')

            doc.add_paragraph()  # spacing

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"questionnaire_{questionnaire.title[:40].replace(' ', '_')}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response


@method_decorator(login_required, name='dispatch')
class ExportPdfView(View):
    def get(self, request, pk):
        from apps.questionnaires.models import Questionnaire
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor, black, grey
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        questionnaire = get_object_or_404(Questionnaire, pk=pk, user=request.user)
        questions = questionnaire.questions.prefetch_related('answers').order_by('order')

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=inch,
            bottomMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()
        brand_blue = HexColor('#1a56DB')
        dark = HexColor('#111827')
        muted = HexColor('#6B7280')
        green = HexColor('#059669')
        red = HexColor('#DC2626')

        title_style = ParagraphStyle('Title', parent=styles['Heading1'],
                                     textColor=brand_blue, fontSize=20, spaceAfter=4)
        subtitle_style = ParagraphStyle('Sub', parent=styles['Normal'],
                                        textColor=muted, fontSize=10, spaceAfter=12)
        q_style = ParagraphStyle('Q', parent=styles['Normal'],
                                 textColor=dark, fontSize=11, fontName='Helvetica-Bold',
                                 spaceBefore=14, spaceAfter=4)
        a_style = ParagraphStyle('A', parent=styles['Normal'],
                                 textColor=dark, fontSize=10, leftIndent=12, spaceAfter=4)
        cite_style = ParagraphStyle('Cite', parent=styles['Normal'],
                                    textColor=muted, fontSize=9, leftIndent=12,
                                    fontName='Helvetica-Oblique')

        story = []
        story.append(Paragraph('LeadFlow AI', title_style))
        story.append(Paragraph(questionnaire.title, ParagraphStyle(
            'H2', parent=styles['Heading2'], fontSize=14, spaceAfter=2)))
        story.append(Paragraph(
            f'Generated {datetime.now().strftime("%B %d, %Y")} &nbsp;|&nbsp; '
            f'{questionnaire.total_questions} questions &nbsp;|&nbsp; '
            f'{questionnaire.answered_questions} answered &nbsp;|&nbsp; '
            f'{questionnaire.not_found_count} not found',
            subtitle_style
        ))
        story.append(HRFlowable(width='100%', thickness=1, color=brand_blue))
        story.append(Spacer(1, 16))

        for question in questions:
            latest = question.answers.order_by('-created_at').first()
            story.append(Paragraph(f'Q{question.order}. {question.text}', q_style))

            if latest:
                color = red if latest.is_not_found else dark
                a_para_style = ParagraphStyle('Ax', parent=a_style, textColor=color)
                story.append(Paragraph(f'<b>Answer:</b> {latest.answer_text}', a_para_style))

                if not latest.is_not_found:
                    conf_pct = int(latest.confidence_score * 100)
                    story.append(Paragraph(
                        f'<i>Confidence: {conf_pct}%</i>',
                        ParagraphStyle('Cf', parent=cite_style, textColor=green)
                    ))

                if latest.citations:
                    sources = ', '.join(c.get('document_title', '') for c in latest.citations)
                    story.append(Paragraph(f'<i>Sources: {sources}</i>', cite_style))
            else:
                story.append(Paragraph('Answer: Not yet generated.', a_style))

        doc.build(story)
        buffer.seek(0)

        filename = f"questionnaire_{questionnaire.title[:40].replace(' ', '_')}.pdf"
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
